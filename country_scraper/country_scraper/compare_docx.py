from itemadapter import ItemAdapter
from scrapy.pipelines.images import FilesPipeline
from country_scraper.items import CountryScraperItem, CSVItem, RtffilesItem
import pandas as pd
from docx.api import Document
from docx.enum.text import WD_COLOR_INDEX
import os
import datetime
from difflib import SequenceMatcher


def len_com(li1, li2):
    if len(li1) > len(li2):
        return True

def add_ind(li1, li2):
    indexes = []
    for index, i in enumerate(li1):
        check = 0
        for j in li2:
            w = SequenceMatcher(None, i[1], j[1])
            if w.ratio() < 0.8:
                check += 1
        if check == len(li2):
            indexes.append(index) 
    for i in indexes:
        li2.insert(i, -1)

def add_ind_table(li1, li2):
    indexes_1 = []
    indexes_2 = []
    for index, i in enumerate(li1):
        check = 0
        for j in li2:
            if i[0] == j[0]:
                w = SequenceMatcher(None, i[3], j[3])
                if w.ratio() < 0.8:
                    check += 1
        if check == len([x for x in li2 if x[0] == i[0]]):
            indexes_1.append(index) 
    
    for i in indexes_1:
        li2.insert(i, -1)
    
    for index, i in enumerate(li2):
        check = 0
        if i != -1:
            for j in li1:
                
                if i[0] == j[0]:
                    w = SequenceMatcher(None, i[3], j[3])
                    if w.ratio() < 0.8:
                        check += 1
            if check == len([x for x in li1 if x[0] == i[0]]):
                indexes_2.append(index) 
    
    for i in indexes_2:
        li1.insert(i, -1)

def compare_doc(item, files):
    file_new = os.getcwd() + '/' + f"docs/{item['country']}" + "/" + files[1][:-4] + '.docx'
    file_old = os.getcwd() + '/' + f"docs/{item['country']}" + "/" + files[0][:-4] + '.docx'

    text = {'Text_old': [], 'Text_new': []}
    text_old = []
    text_new = []
    t_new = []
    t_old = []

    document = Document(file_new)
    document_old = Document(file_old)

    length_para = max(len(document.paragraphs),len(document_old.paragraphs))

    for i in range(length_para):
        try:
            text1 = document.paragraphs[i].text
            text0 = document_old.paragraphs[i].text
        except IndexError:
            pass
        t_new.append((i, text1))
        t_old.append((i, text0))

    diff = lambda l1,l2: [x for x in l1 if x[1] not in [y[1] for y in l2]]
    diff_new = diff(t_new,t_old)
    dif_old = diff(t_old, t_new)

    if len_com(diff_new, dif_old):
        add_ind(diff_new, dif_old)
    else:
        add_ind(dif_old, diff_new)

    for i in diff_new:
        if i == -1:
            text_new.append('')
        else:
            text_new.append(document.paragraphs[i[0]].text)
            for run in document.paragraphs[i[0]].runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for i in dif_old:
        if i == -1:
            text_old.append('')
        else:
            text_old.append(document_old.paragraphs[i[0]].text)
            for run in document_old.paragraphs[i[0]].runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for i in range(len(text_new)):
        text['Text_new'].append(text_new[i])
        text['Text_old'].append(text_old[i])


    cont_table = []
    cont_table_old = []
    for ind1, i in enumerate(document.tables):
        for ind2, j in enumerate(i.rows):
            for ind3, c in enumerate(j.cells):
                cont_table.append((ind1, ind2, ind3, c.text))

    for ind1, i in enumerate(document_old.tables):
        for ind2, j in enumerate(i.rows):
            for ind3, c in enumerate(j.cells):
                cont_table_old.append((ind1, ind2, ind3, c.text))

    diff_table = lambda l1,l2: [x for index, x in enumerate(l1) if x[3] not in [y[3] for y in l2 if y[0] == x[0]]]

    dif_table_new = diff_table(cont_table, cont_table_old)
    dif_table_old = diff_table(cont_table_old, cont_table)

    add_ind_table(dif_table_new, dif_table_old)

    for i in dif_table_old:
        if i == -1:
            text['Text_old'].append('')
        else:
            text['Text_old'].append(i[3])
    
    for i in dif_table_new:
        if i == -1:
            text['Text_new'].append('')
        else:
            text['Text_new'].append(i[3])

    for i in dif_table_new:
        if i != -1:
            for index, j in enumerate(document.tables[i[0]].rows[i[1]].cells):
                if index == i[2]:
                    for p in j.paragraphs:
                        for run in p.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for i in dif_table_old:
        if i != -1:
            for index, j in enumerate(document_old.tables[i[0]].rows[i[1]].cells):
                if index == i[2]:
                    for p in j.paragraphs:
                        for run in p.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    df = pd.DataFrame(text)
    df.to_csv(os.getcwd() + '/' + f"docs/{item['country']}" + "/" + f'ver_diff_{files[0][:4]}_{str(datetime.datetime.now())[:10]}.csv', sep=',')
    document.save(os.getcwd() + '/' + f"docs/{item['country']}" + "/" + files[1][:-4] + '.docx')
    document_old.save(os.getcwd() + '/' + f"docs/{item['country']}" + "/" + files[0][:-4] + '.docx')
