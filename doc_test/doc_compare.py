import pandas as pd
from docx.api import Document
from docx.enum.text import WD_COLOR_INDEX
import os
from difflib import SequenceMatcher
#w = SequenceMatcher(None, i, j)
			#if w.ratio() > 0.95:


# file = 'c:/job/python/TAR_EMEAR/doc_test/Text1.docx'
# file = 'c:/job/python/TAR_EMEAR/doc_test/1960.03.15_2022-04-01_1.docx'
file = 'c:/job/python/TAR_EMEAR/doc_test/Text1.docx'
# file_old = 'c:/job/python/TAR_EMEAR/doc_test/1960.03.15_2022-04-01.docx'
file_old = 'c:/job/python/TAR_EMEAR/doc_test/Text.docx'


text = {'Text_old': [], 'Text_new': []}
text_old = []
text_new = []
t_new = []
t_old = []

document = Document(file)
document_old = Document(file_old)


length_para = max(len(document.paragraphs),len(document_old.paragraphs))

for i in range(length_para):
    try:
        text1 = document.paragraphs[i].text
        text0 = document_old.paragraphs[i].text
        
    except IndexError:
        pass
    t_new.append((i, text1))
    t_old.append((i, text0))

diff = lambda l1,l2: [x for x in l1 if x[1] not in [y[1] for y in l2]]
qq = diff(t_new,t_old)
qq1 = diff(t_old, t_new)
# print(qq)
# print(qq1)

def len_com(li1, li2):
    if len(li1) >= len(li2):
        return True

def add_ind(li1, li2):
    indexes = []
    for index, i in enumerate(li1):
        check = 0
        for j in li2:
            w = SequenceMatcher(None, i[1], j[1])
            if w.ratio() < 0.8:
                check += 1
        if check == len(li2):
            indexes.append(index) 
    for i in indexes:
        li2.insert(i, -1)
# print(len_com(qq1, qq))

a = [64, 214, 303, 314]
b = [212, 301, 312, 393]    
# print(len_com(a,b))

if len_com(qq, qq1):
    add_ind(qq, qq1)
else:
    add_ind(qq1, qq)

indexes = []
# for index, i in enumerate(qq):
    # check = 0
    # for j in qq1:
        # if abs(i-j) > 2:
            # check += 1
    # if check == len(qq1):
        # indexes.append(index) 
        
# print(indexes)    
# for i in indexes:
    # li2.insert(i, -1)

# tex_text = [i for i in t_new if i[0] in qq]
# tex_text_old = [i for i in t_old if i[0] in qq1]
# for index, i in enumerate(tex_text):
    # check = 0
    # for  j in tex_text_old:
        # w = SequenceMatcher(None, i[1], j[1])
        # if w.ratio() < 0.8:
            # print(w.ratio())
            # check += 1
            # print(i[1])
    # print(check)
    # if check == len(tex_text_old):
        # indexes.append(index) 
        
# print(indexes)
# for i in indexes:
    # tex_text_old.insert(i, -1)

# print(qq)
# print(qq1)  
    
# print(tex_text)
# print(tex_text_old)



# print(qq1)

for i in qq:
    if i == -1:
        text_new.append('')
    else:
        text_new.append(document.paragraphs[i[0]].text)
        for run in document.paragraphs[i[0]].runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
for i in qq1:
    if i == -1:
        text_old.append('')
    else:
        text_old.append(document_old.paragraphs[i[0]].text)
        for run in document_old.paragraphs[i[0]].runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

for i in range(len(text_new)):
    text['Text_new'].append(text_new[i])
    text['Text_old'].append(text_old[i])
      
# print(text)







 
                
    # for index1, j in enumerate(document_old.paragraphs):
        # if index == index1:
        # w = SequenceMatcher(None, i.text.strip(), j.text.strip())
        
            # if i.text != j.text:
        # if w.ratio() < 0.9:
            # print(i.text)
            
            # text['Text_new'].append(i.text)
            # text['Text_old'].append(j.text)
            # for run in i.runs:
                # run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            # for run in j.runs:
                # run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_ind_table(li1, li2):
    indexes_1 = []
    indexes_2 = []
    for index, i in enumerate(li1):
        check = 0
        for j in li2:
            if i[0] == j[0]:
                w = SequenceMatcher(None, i[3], j[3])
                if w.ratio() < 0.8:
                    check += 1
        # print(len([x for x in li2 if x[0] == i[0]]))
        if check == len([x for x in li2 if x[0] == i[0]]):
            indexes_1.append(index) 
    
    for i in indexes_1:
        li2.insert(i, -1)
    
    for index, i in enumerate(li2):
        check = 0
        if i != -1:
            for j in li1:
                
                if i[0] == j[0]:
                    w = SequenceMatcher(None, i[3], j[3])
                    if w.ratio() < 0.8:
                        check += 1
            if check == len([x for x in li1 if x[0] == i[0]]):
                indexes_2.append(index) 
    
    for i in indexes_2:
        li1.insert(i, -1)





cont_table = []
cont_table_old = []
for ind1, i in enumerate(document.tables):
    for ind2, j in enumerate(i.rows):
        for ind3, c in enumerate(j.cells):
            cont_table.append((ind1, ind2, ind3, c.text))

for ind1, i in enumerate(document_old.tables):
    for ind2, j in enumerate(i.rows):
        for ind3, c in enumerate(j.cells):
            cont_table_old.append((ind1, ind2, ind3, c.text))
            
            
            
# length_table = max(len(document.paragraphs),len(document_old.paragraphs))

# for i in range(length_para):
    # try:
        # text1 = document.paragraphs[i].text
        # text0 = document_old.paragraphs[i].text
        
    # except IndexError:
        # pass
    # t_new.append((i, text1))
    # t_old.append((i, text0))            



# print(cont_table)
# print(cont_table_old)
diff_table = lambda l1,l2: [x for index, x in enumerate(l1) if x[3] not in [y[3] for y in l2 if y[0] == x[0]]]
# diff_table = lambda l1,l2: [x for index, x in enumerate(l1) if x[3] not in [y[3] for y in l2] and x[0] == l2[index][0]]


dif_table_new = diff_table(cont_table, cont_table_old)
dif_table_old = diff_table(cont_table_old, cont_table)

add_ind_table(dif_table_new, dif_table_old)
# print(len(dif_table_new))
# print(len(dif_table_old))
# for i in dif_table_new:
    # print(i)
# print('------------------')
# for i in dif_table_old:
    # print(i)
    
# if len_com(dif_table_new, dif_table_old):

# else:
    # add_ind_table(dif_table_old, dif_table_new)

# print(dif_table_old)
# print(dif_table_new)
# print(set([x for x in cont_table for y in cont_table_old if x[3] not in [y[3] for y in cont_table_old] and x[0] == y[0]]))
# print(set([x for x in cont_table_old for y in cont_table if x[3] not in [y[3] for y in cont_table if y[0] == x[0]]]))
for i in dif_table_old:
    if i == -1:
        text['Text_old'].append('')
    else:
        text['Text_old'].append(i[3])
    

for i in dif_table_new:
    if i == -1:
        text['Text_new'].append('')
    else:
        text['Text_new'].append(i[3])


# print(text)
# li = [y[3] for y in cont_table_old]
# print(li)

# for index, x in enumerate(cont_table):
    # print(x[0], x[3])
    # print(cont_table_old[index][0], cont_table_old[index][3])
    # print('-----')


for i in dif_table_new:
    if i != -1:
        for index, j in enumerate(document.tables[i[0]].rows[i[1]].cells):
            if index == i[2]:
                for p in j.paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

for i in dif_table_old:
    if i != -1:
        for index, j in enumerate(document_old.tables[i[0]].rows[i[1]].cells):
            if index == i[2]:
                for p in j.paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW






# print(text)
df = pd.DataFrame(text)
df.to_csv('c:/job/python/TAR_EMEAR/doc_test/ver_diff.csv', sep=',')
# print(df)
# document.save('c:/job/python/TAR_EMEAR/doc_test/1960.03.15_2022-04-01_1.docx')
document.save('c:/job/python/TAR_EMEAR/doc_test/Text1.docx')
# document_old.save('c:/job/python/TAR_EMEAR/doc_test/1960.03.15_2022-04-01.docx')
document_old.save('c:/job/python/TAR_EMEAR/doc_test/Text.docx')


def compare_doc1(item, files):
    file_new = os.getcwd() + '/' + f"docs/{item['country']}" + "/" + files[1][:-4] + '.docx'
    file_old = os.getcwd() + '/' + f"docs/{item['country']}" + "/" + files[0][:-4] + '.docx'

    text = {'Text_old': [], 'Text_new': []}

    document = Document(file_new)
    document_old = Document(file_old)

    for index, i in enumerate(document.paragraphs):
        for index1, j in enumerate(document_old.paragraphs):
            if index == index1:
                if i.text != j.text:
                    text['Text_new'].append(i.text)
                    text['Text_old'].append(j.text)
                    for run in i.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    for run in j.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for ind1, i in enumerate(document.tables):
        for ind2, j in enumerate(document_old.tables):
            if ind1 == ind2:
                for r_ind1, row in enumerate(i.rows):
                    for r_ind2, row_old in enumerate(j.rows):
                        if r_ind1 == r_ind2:
                            for c_ind1, cell in enumerate(row.cells):
                                for c_ind2, cell_old in enumerate(row_old.cells):
                                    if c_ind1 == c_ind2:
                                        if cell.text != cell_old.text:
                                            for p in cell.paragraphs:
                                                text['Text_new'].append(p.text)
                                                for run in p.runs:
                                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                            for p in cell_old.paragraphs:
                                                text['Text_old'].append(p.text)
                                                for run in p.runs:
                                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW       

    df = pd.DataFrame(text)
    df.to_csv(os.getcwd() + '/' + f"docs/{item['country']}" + "/" + f'ver_diff_{files[0][:4]}_{str(datetime.datetime.now())[:10]}.csv', sep=',')
    document.save(os.getcwd() + '/' + f"docs/{item['country']}" + "/" + files[1][:-4] + '.docx')
    document_old.save(os.getcwd() + '/' + f"docs/{item['country']}" + "/" + files[0][:-4] + '.docx')